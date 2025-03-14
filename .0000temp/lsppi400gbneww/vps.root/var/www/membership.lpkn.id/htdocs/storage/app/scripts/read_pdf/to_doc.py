import pdfplumber
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt
import os
import shutil
import sys
import random
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT  # Untuk mengatur alignment teks
import json  # Tambahkan modul JSON


# Nama file PDF dan output DOCX

# dynamic pdf
# Pastikan ada argumen file PDF
if len(sys.argv) < 2:
    print(json.dumps({"error": "No PDF file provided"}))
    sys.exit(1)

pdf_path = os.path.abspath(sys.argv[1])  # Ambil path file dari argumen

# Cek apakah file PDF ada
if not os.path.exists(pdf_path):
    print(json.dumps({"error": f"File not found: {pdf_path}"}))
    sys.exit(1)
# end dynamic pdf

# pdf_path = os.path.abspath("uploaded_files/pdf_presensi_lkpp/file-contoh.pdf")
# pdf_path = "file.pdf"

folder_path = "uploaded_files/data_peserta"
# Cek apakah folder sudah ada, jika tidak maka buat
if not os.path.exists(folder_path):
    os.makedirs(folder_path)

# Buat nama file secara acak
randFolder = random.randrange(1, 999)
filees = f"data_peserta-{randFolder}.docx"
docx_path = os.path.join(folder_path, filees)  # Gabungkan path dengan aman

# Folder untuk menyimpan gambar
image_folder = f"extracted_images/{randFolder}"
if not os.path.exists(image_folder):
    os.makedirs(image_folder)

# Fungsi untuk menyimpan gambar dari PDF lama
def save_image(image_bytes, page_num, img_index):
    img_path = os.path.join(image_folder, f"image_page{page_num+1}_img{img_index+1}.png")
    with open(img_path, "wb") as f:
        f.write(image_bytes)
    return img_path

# Step 1: Ekstrak data (Nama, Instansi) dari tabel PDF
data_list = []
with pdfplumber.open(pdf_path) as pdf:
    for page in pdf.pages:
        tables = page.extract_tables()
        for table in tables:
            for row in table:
                if len(row) >= 2:  # Pastikan ada data Nama & Instansi
                    noo_urut = row[0].strip()
                    no_ujian = row[1].strip()
                    nama = row[3].strip().replace("\n", " ")
                    nik  = row[4]
                    nip  = row[5]
                    email  = row[6]
                    instansi = row[7].strip().replace("\n", " ")
                    if instansi.endswith(','):
                        instansi = instansi[:-1]
                    if nama != 'Nama':
                        data_list.append({"noo_urut":noo_urut, "no_ujian":no_ujian, "nama": nama, "nama":nama, 
                            "nik":nik, "nip":nip, "email":email, "instansi": instansi, "foto": None})                    

# Step 2: Ekstrak gambar dari PDF
pdf_document = fitz.open(pdf_path)
image_list = []
for page_num in range(len(pdf_document)):
    for img_index, img in enumerate(pdf_document[page_num].get_images(full=True)):
        xref = img[0]
        base_image = pdf_document.extract_image(xref)
        image_bytes = base_image["image"]

        # Simpan gambar
        img_path = save_image(image_bytes, page_num, img_index)
        image_list.append(img_path)

# **Asosiasikan gambar dengan data Nama & Instansi**
for i in range(min(len(data_list), len(image_list))):
    data_list[i]["foto"] = image_list[i]

# Step 3: Buat file DOCX dengan Tabel
doc = Document()

# Buat tabel dengan 3 kolom (Nama, Instansi, Foto)
table = doc.add_table(rows=1, cols=3)
table.style = "Table Grid"

# Header tabel
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "No"
hdr_cells[1].text = "Data Peserta"
hdr_cells[2].text = "Foto"

# Tambahkan data ke tabel
for data in data_list:
    row_cells = table.add_row().cells
    row_cells[0].text = data["noo_urut"]
    p2 = row_cells[0].paragraphs[0]
    p2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # **Atur gambar ke tengah**

    row_cells[1].text = ""
    p = row_cells[1].paragraphs[0]
    p.add_run("\n")
    p.add_run("Nomor Ujian : " + data['no_ujian']).bold = True
    p.add_run("\n")
    p.add_run("Nama : " + data['nama'])
    p.add_run("\n")
    p.add_run("NIK : " + data['nik'])
    p.add_run("\n")
    p.add_run("NIP : " + data['nip'])
    p.add_run("\n")
    p.add_run("Asal Instansi : " + data['instansi'])
    p.add_run("\n")
    p.add_run("\n")

    # Atur font size kolom "No"
    for paragraph in row_cells[1].paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(16)  # Misalnya: 16pt

    if data["foto"]:
        # Sisipkan gambar di kolom foto
        paragraph = row_cells[2].paragraphs[0]
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER  # **Atur gambar ke tengah**
        run = paragraph.add_run("\n")
        run.add_picture(data["foto"], width=Inches(1.2), height=Inches(1.8))  # Ukuran gambar

#hapus folder gambar
if os.path.exists(image_folder):
    shutil.rmtree(image_folder)

# Simpan hasil ke file DOCX
doc.save(docx_path)
response = {
    "status": "success",
    "data": data_list,
    "file_datapeserta": docx_path
}
response_json = json.dumps(response, indent=4)
print(response_json)  # ✅ Response JSON bisa digunakan dalam API